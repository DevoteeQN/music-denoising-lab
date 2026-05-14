from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "report" / "music_denoising_report.md"
DOCX_PATH = ROOT / "report" / "music_denoising_report.docx"


def set_font(paragraph, size: int = 11, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold


def add_code_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)


def is_table_start(lines: list[str], idx: int) -> bool:
    return (
        idx + 1 < len(lines)
        and lines[idx].strip().startswith("|")
        and lines[idx + 1].strip().startswith("|")
        and set(lines[idx + 1].strip().replace("|", "").replace(":", "").replace(" ", "")) <= {"-"}
    )


def parse_table(lines: list[str], idx: int):
    table_lines = []
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        table_lines.append(lines[idx].strip())
        idx += 1
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-+:?", cell.replace(" ", "")) for cell in cells):
            continue
        rows.append(cells)
    return rows, idx


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            table.cell(r_idx, c_idx).text = cell
            for paragraph in table.cell(r_idx, c_idx).paragraphs:
                set_font(paragraph, size=10, bold=(r_idx == 0))


def add_appendix_images(doc: Document) -> None:
    figures = [
        ("多方法波形对比", ROOT / "outputs" / "figures" / "waveform_comparison_all.png"),
        ("多方法频谱图对比", ROOT / "outputs" / "figures" / "spectrogram_comparison_all.png"),
        ("U-Net loss 曲线", ROOT / "outputs" / "figures" / "loss_curve.png"),
        ("TCN loss 曲线", ROOT / "outputs" / "figures" / "tcn_loss_curve.png"),
        ("CNN loss 曲线", ROOT / "outputs" / "figures" / "cnn_loss_curve.png"),
        ("Tiny DiT-style loss 曲线", ROOT / "outputs" / "figures" / "dit_loss_curve.png"),
    ]
    doc.add_heading("附录：关键实验图", level=1)
    for title, path in figures:
        if not path.exists():
            continue
        p = doc.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p, size=11, bold=True)
        doc.add_picture(str(path), width=Inches(6.2))


def add_markdown_image(doc: Document, alt: str, image_ref: str) -> None:
    image_path = Path(image_ref)
    if not image_path.is_absolute():
        image_path = MD_PATH.parent / image_path
    if not image_path.exists():
        p = doc.add_paragraph(f"[图片缺失：{alt} -> {image_ref}]")
        set_font(p, size=10)
        return
    p = doc.add_paragraph(alt)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p, size=10, bold=True)
    doc.add_picture(str(image_path), width=Inches(6.2))


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_buffer: list[str] = []
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code_paragraph(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_buffer.append(line)
            idx += 1
            continue
        if not stripped:
            idx += 1
            continue
        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        if image_match:
            add_markdown_image(doc, image_match.group(1), image_match.group(2))
            idx += 1
            continue
        if is_table_start(lines, idx):
            rows, idx = parse_table(lines, idx)
            add_markdown_table(doc, rows)
            continue
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            idx += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            idx += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            idx += 1
            continue
        if stripped.startswith(">"):
            p = doc.add_paragraph(stripped.lstrip("> ").strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p, size=10)
            idx += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
            set_font(p)
            idx += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            set_font(p)
            idx += 1
            continue
        p = doc.add_paragraph(stripped)
        set_font(p)
        idx += 1

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
